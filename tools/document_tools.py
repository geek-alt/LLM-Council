"""
Document Ingestion Tools — Multi-modal document parsing for PDFs, DOCX, GitHub repos, and more.
Supports intelligent chunking, citation tracking, and metadata extraction.
"""

import logging
import hashlib
import re
from pathlib import Path
from typing import Optional, List, Dict, Any
from dataclasses import dataclass, field
from urllib.parse import urlparse

logger = logging.getLogger("council.documents")


@dataclass
class DocumentChunk:
    """A single chunk of a parsed document with metadata."""
    content: str
    doc_id: str
    doc_title: str
    doc_source: str  # file path or URL
    chunk_index: int
    total_chunks: int
    start_page: Optional[int] = None
    end_page: Optional[int] = None
    metadata: Dict = field(default_factory=dict)
    
    def to_citation(self) -> str:
        """Generate a citation string for this chunk."""
        parts = [f"[{self.doc_title}"]
        if self.start_page is not None:
            if self.end_page and self.end_page != self.start_page:
                parts.append(f"pp. {self.start_page}-{self.end_page}")
            else:
                parts.append(f"p. {self.start_page}")
        parts.append(f"— Chunk {self.chunk_index + 1}/{self.total_chunks}]")
        return " ".join(parts)


@dataclass
class ParsedDocument:
    """A fully parsed document with chunks and metadata."""
    doc_id: str
    title: str
    source: str  # file path or URL
    source_type: str  # pdf, docx, github, url, text
    chunks: List[DocumentChunk]
    metadata: dict = field(default_factory=dict)
    full_text: str = ""
    
    def to_dict(self) -> dict:
        return {
            "doc_id": self.doc_id,
            "title": self.title,
            "source": self.source,
            "source_type": self.source_type,
            "chunk_count": len(self.chunks),
            "metadata": self.metadata,
            "full_text": self.full_text[:2000] + "..." if len(self.full_text) > 2000 else self.full_text,
        }


class DocumentParser:
    """Base class for document parsers."""
    
    def parse(self, source: str) -> ParsedDocument:
        raise NotImplementedError
    
    def _generate_doc_id(self, content: str, source: str) -> str:
        """Generate a unique document ID based on content hash."""
        hash_input = f"{source}:{content[:500]}"
        return hashlib.md5(hash_input.encode()).hexdigest()[:12]
    
    def _chunk_text(self, text: str, max_chunk_size: int = 1500, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks with sentence-aware boundaries."""
        if len(text) <= max_chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + max_chunk_size
            
            if end >= len(text):
                chunks.append(text[start:].strip())
                break
            
            # Try to break at sentence boundary
            last_period = text.rfind(".", start, end)
            last_newline = text.rfind("\n", start, end)
            
            break_point = max(last_period, last_newline)
            if break_point > start + max_chunk_size // 2:
                end = break_point + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
        
        return chunks


class PDFParser(DocumentParser):
    """PDF document parser using PyMuPDF (fitz)."""
    
    def __init__(self, max_chunk_size: int = 1500, chunk_overlap: int = 200):
        self.max_chunk_size = max_chunk_size
        self.chunk_overlap = chunk_overlap
    
    def parse(self, source: str) -> ParsedDocument:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.error("PyMuPDF not installed. Run: pip install PyMuPDF")
            raise ImportError("PyMuPDF required for PDF parsing")
        
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"PDF not found: {source}")
        
        logger.info(f"Parsing PDF: {path.name}")
        
        try:
            doc = fitz.open(path)
            full_text = ""
            page_texts = []
            
            for page in doc:
                page_text = page.get_text()
                page_texts.append(page_text)
                full_text += page_text + "\n\n"
            
            doc.close()
            
            title = path.stem
            doc_id = self._generate_doc_id(full_text, str(path))
            chunks = self._chunk_text(full_text, self.max_chunk_size, self.chunk_overlap)
            
            doc_chunks = []
            for idx, chunk in enumerate(chunks):
                # Try to determine page range for this chunk
                char_pos = full_text.find(chunk[:50])
                start_page = None
                end_page = None
                
                if char_pos >= 0:
                    cum_len = 0
                    for p_idx, p_text in enumerate(page_texts):
                        cum_len += len(p_text)
                        if start_page is None and char_pos < cum_len:
                            start_page = p_idx + 1
                        if char_pos + len(chunk) < cum_len:
                            end_page = p_idx + 1
                            break
                
                doc_chunks.append(DocumentChunk(
                    content=chunk,
                    doc_id=doc_id,
                    doc_title=title,
                    doc_source=str(path),
                    chunk_index=idx,
                    total_chunks=len(chunks),
                    start_page=start_page,
                    end_page=end_page or start_page,
                    metadata={"source_type": "pdf", "file_path": str(path)},
                ))
            
            return ParsedDocument(
                doc_id=doc_id,
                title=title,
                source=str(path),
                source_type="pdf",
                chunks=doc_chunks,
                metadata={"page_count": len(page_texts), "file_size": path.stat().st_size},
                full_text=full_text,
            )
            
        except Exception as e:
            logger.error(f"Failed to parse PDF {source}: {e}")
            raise


class DOCXParser(DocumentParser):
    """DOCX document parser using python-docx."""
    
    def __init__(self, max_chunk_size: int = 1500, chunk_overlap: int = 200):
        self.max_chunk_size = max_chunk_size
        self.chunk_overlap = chunk_overlap
    
    def parse(self, source: str) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise ImportError("python-docx required for DOCX parsing")
        
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"DOCX not found: {source}")
        
        logger.info(f"Parsing DOCX: {path.name}")
        
        try:
            doc = Document(path)
            full_text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            title = path.stem
            doc_id = self._generate_doc_id(full_text, str(path))
            chunks = self._chunk_text(full_text, self.max_chunk_size, self.chunk_overlap)
            
            doc_chunks = []
            for idx, chunk in enumerate(chunks):
                doc_chunks.append(DocumentChunk(
                    content=chunk,
                    doc_id=doc_id,
                    doc_title=title,
                    doc_source=str(path),
                    chunk_index=idx,
                    total_chunks=len(chunks),
                    metadata={"source_type": "docx", "file_path": str(path)},
                ))
            
            return ParsedDocument(
                doc_id=doc_id,
                title=title,
                source=str(path),
                source_type="docx",
                chunks=doc_chunks,
                metadata={"paragraph_count": len(doc.paragraphs), "file_size": path.stat().st_size},
                full_text=full_text,
            )
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX {source}: {e}")
            raise


class TextParser(DocumentParser):
    """Plain text file parser."""
    
    def __init__(self, max_chunk_size: int = 1500, chunk_overlap: int = 200):
        self.max_chunk_size = max_chunk_size
        self.chunk_overlap = chunk_overlap
    
    def parse(self, source: str) -> ParsedDocument:
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"Text file not found: {source}")
        
        logger.info(f"Parsing text file: {path.name}")
        
        try:
            full_text = path.read_text(encoding="utf-8", errors="ignore")
            title = path.stem
            doc_id = self._generate_doc_id(full_text, str(path))
            chunks = self._chunk_text(full_text, self.max_chunk_size, self.chunk_overlap)
            
            doc_chunks = []
            for idx, chunk in enumerate(chunks):
                doc_chunks.append(DocumentChunk(
                    content=chunk,
                    doc_id=doc_id,
                    doc_title=title,
                    doc_source=str(path),
                    chunk_index=idx,
                    total_chunks=len(chunks),
                    metadata={"source_type": "text", "file_path": str(path)},
                ))
            
            return ParsedDocument(
                doc_id=doc_id,
                title=title,
                source=str(path),
                source_type="text",
                chunks=doc_chunks,
                metadata={"file_size": path.stat().st_size, "char_count": len(full_text)},
                full_text=full_text,
            )
            
        except Exception as e:
            logger.error(f"Failed to parse text file {source}: {e}")
            raise


class GitHubRepoParser(DocumentParser):
    """GitHub repository parser that fetches key files."""
    
    def __init__(self, max_chunk_size: int = 1500, chunk_overlap: int = 200):
        self.max_chunk_size = max_chunk_size
        self.chunk_overlap = chunk_overlap
        self.key_files = [
            "README.md", "README.txt", "README",
            "package.json", "requirements.txt", "pyproject.toml",
            "Cargo.toml", "go.mod", "pom.xml", "build.gradle",
            "setup.py", "Makefile", "Dockerfile",
            ".github/ISSUE_TEMPLATE.md", "CONTRIBUTING.md", "LICENSE",
        ]
    
    def _parse_github_url(self, url: str) -> tuple[str, str]:
        """Extract owner and repo name from GitHub URL."""
        parsed = urlparse(url)
        path_parts = parsed.path.strip("/").split("/")
        
        if len(path_parts) >= 2:
            owner = path_parts[0]
            repo = path_parts[1].replace(".git", "")
            return owner, repo
        
        raise ValueError(f"Invalid GitHub URL: {url}")
    
    def parse(self, source: str) -> ParsedDocument:
        try:
            import requests
        except ImportError:
            logger.error("requests not installed")
            raise ImportError("requests required for GitHub parsing")
        
        owner, repo = self._parse_github_url(source)
        base_api = f"https://api.github.com/repos/{owner}/{repo}"
        
        logger.info(f"Parsing GitHub repo: {owner}/{repo}")
        
        full_text = f"# Repository: {owner}/{repo}\n\n"
        fetched_files = []
        
        # Fetch README first
        readme_content = self._fetch_file(base_api, "README.md")
        if readme_content:
            full_text += f"## README\n{readme_content}\n\n"
            fetched_files.append("README.md")
        
        # Fetch other key files
        for filename in self.key_files[1:]:
            content = self._fetch_file(base_api, filename)
            if content:
                full_text += f"## {filename}\n{content}\n\n"
                fetched_files.append(filename)
        
        if not fetched_files:
            full_text += "No key documentation files found.\n"
        
        title = f"{owner}/{repo}"
        doc_id = self._generate_doc_id(full_text, source)
        chunks = self._chunk_text(full_text, self.max_chunk_size, self.chunk_overlap)
        
        doc_chunks = []
        for idx, chunk in enumerate(chunks):
            doc_chunks.append(DocumentChunk(
                content=chunk,
                doc_id=doc_id,
                doc_title=title,
                doc_source=source,
                chunk_index=idx,
                total_chunks=len(chunks),
                metadata={"source_type": "github", "owner": owner, "repo": repo, "files": fetched_files},
            ))
        
        return ParsedDocument(
            doc_id=doc_id,
            title=title,
            source=source,
            source_type="github",
            chunks=doc_chunks,
            metadata={"owner": owner, "repo": repo, "files_fetched": fetched_files},
            full_text=full_text,
        )
    
    def _fetch_file(self, base_api: str, filepath: str) -> Optional[str]:
        """Fetch a file from GitHub API."""
        try:
            import requests
            
            # Get file info
            resp = requests.get(f"{base_api}/contents/{filepath}", timeout=10)
            if resp.status_code != 200:
                return None
            
            data = resp.json()
            
            if isinstance(data, list):
                return None  # Directory, not file
            
            download_url = data.get("download_url")
            if not download_url:
                return None
            
            # Fetch raw content
            content_resp = requests.get(download_url, timeout=10)
            if content_resp.status_code == 200:
                return content_resp.text[:10000]  # Limit file size
            
        except Exception as e:
            logger.debug(f"Failed to fetch {filepath}: {e}")
        
        return None


class URLParser(DocumentParser):
    """Web page parser using requests + BeautifulSoup."""
    
    def __init__(self, max_chunk_size: int = 1500, chunk_overlap: int = 200):
        self.max_chunk_size = max_chunk_size
        self.chunk_overlap = chunk_overlap
    
    def parse(self, source: str) -> ParsedDocument:
        try:
            import requests
            from bs4 import BeautifulSoup
        except ImportError:
            logger.error("Missing dependencies. Run: pip install beautifulsoup4")
            raise ImportError("beautifulsoup4 required for URL parsing")
        
        logger.info(f"Parsing URL: {source}")
        
        try:
            headers = {"User-Agent": "Mozilla/5.0 (LLM Council Document Parser)"}
            resp = requests.get(source, headers=headers, timeout=30)
            resp.raise_for_status()
            
            soup = BeautifulSoup(resp.text, "html.parser")
            
            # Extract title
            title_tag = soup.find("title")
            title = title_tag.get_text(strip=True) if title_tag else urlparse(source).netloc
            
            # Extract main content (remove scripts, styles, navs)
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            
            # Get text from meaningful elements
            text_elements = []
            for tag in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "blockquote", "article"]):
                text = tag.get_text(strip=True)
                if text and len(text) > 20:
                    text_elements.append(text)
            
            full_text = "\n\n".join(text_elements)
            
            if not full_text.strip():
                full_text = soup.get_text(separator="\n", strip=True)
            
            doc_id = self._generate_doc_id(full_text, source)
            chunks = self._chunk_text(full_text, self.max_chunk_size, self.chunk_overlap)
            
            doc_chunks = []
            for idx, chunk in enumerate(chunks):
                doc_chunks.append(DocumentChunk(
                    content=chunk,
                    doc_id=doc_id,
                    doc_title=title,
                    doc_source=source,
                    chunk_index=idx,
                    total_chunks=len(chunks),
                    metadata={"source_type": "url", "domain": urlparse(source).netloc},
                ))
            
            return ParsedDocument(
                doc_id=doc_id,
                title=title,
                source=source,
                source_type="url",
                chunks=doc_chunks,
                metadata={"domain": urlparse(source).netloc, "word_count": len(full_text.split())},
                full_text=full_text,
            )
            
        except Exception as e:
            logger.error(f"Failed to parse URL {source}: {e}")
            raise


class DocumentIngestionEngine:
    """
    Main engine for multi-modal document ingestion.
    Automatically detects document type and routes to appropriate parser.
    """
    
    def __init__(
        self,
        max_chunk_size: int = 1500,
        chunk_overlap: int = 200,
        auto_detect: bool = True,
    ):
        self.max_chunk_size = max_chunk_size
        self.chunk_overlap = chunk_overlap
        self.auto_detect = auto_detect
        
        self.parsers = {
            "pdf": PDFParser(max_chunk_size, chunk_overlap),
            "docx": DOCXParser(max_chunk_size, chunk_overlap),
            "text": TextParser(max_chunk_size, chunk_overlap),
            "github": GitHubRepoParser(max_chunk_size, chunk_overlap),
            "url": URLParser(max_chunk_size, chunk_overlap),
        }
        
        self.parsed_documents: Dict[str, ParsedDocument] = {}
    
    def _detect_type(self, source: str) -> str:
        """Auto-detect document type from source."""
        if source.startswith("http://") or source.startswith("https://"):
            if "github.com" in source:
                return "github"
            return "url"
        
        path = Path(source)
        suffix = path.suffix.lower()
        
        if suffix == ".pdf":
            return "pdf"
        elif suffix == ".docx":
            return "docx"
        elif suffix in [".txt", ".md", ".rst", ".tex"]:
            return "text"
        
        # Default to text for unknown extensions
        return "text"
    
    def ingest(self, source: str, doc_type: Optional[str] = None) -> ParsedDocument:
        """
        Ingest a document from various sources.
        
        Args:
            source: File path, URL, or GitHub repo URL
            doc_type: Optional override for auto-detection (pdf, docx, text, github, url)
        
        Returns:
            ParsedDocument with chunks and metadata
        """
        if doc_type is None and self.auto_detect:
            doc_type = self._detect_type(source)
        elif doc_type is None:
            raise ValueError("Cannot determine document type. Set doc_type parameter.")
        
        if doc_type not in self.parsers:
            raise ValueError(f"Unsupported document type: {doc_type}")
        
        parser = self.parsers[doc_type]
        parsed = parser.parse(source)
        
        self.parsed_documents[parsed.doc_id] = parsed
        logger.info(f"Ingested {doc_type} document: {parsed.title} ({len(parsed.chunks)} chunks)")
        
        return parsed
    
    def ingest_batch(self, sources: List[tuple[str, Optional[str]]]) -> List[ParsedDocument]:
        """Ingest multiple documents."""
        results = []
        for source, doc_type in sources:
            try:
                parsed = self.ingest(source, doc_type)
                results.append(parsed)
            except Exception as e:
                logger.error(f"Failed to ingest {source}: {e}")
        return results
    
    def get_all_chunks(self) -> List[DocumentChunk]:
        """Get all chunks from all ingested documents."""
        all_chunks = []
        for doc in self.parsed_documents.values():
            all_chunks.extend(doc.chunks)
        return all_chunks
    
    def get_chunks_by_doc_id(self, doc_id: str) -> List[DocumentChunk]:
        """Get chunks for a specific document."""
        doc = self.parsed_documents.get(doc_id)
        return doc.chunks if doc else []
    
    def format_chunks_for_context(self, max_chunks: int = 10, max_total_chars: int = 8000) -> str:
        """Format chunks for injection into LLM context with citations."""
        all_chunks = self.get_all_chunks()
        
        if not all_chunks:
            return "No documents ingested."
        
        # Sort by chunk index within each doc
        all_chunks.sort(key=lambda c: (c.doc_id, c.chunk_index))
        
        formatted = []
        total_chars = 0
        
        for chunk in all_chunks[:max_chunks]:
            citation = chunk.to_citation()
            entry = f"{citation}\n{chunk.content}"
            
            if total_chars + len(entry) > max_total_chars:
                break
            
            formatted.append(entry)
            total_chars += len(entry)
        
        header = f"=== INGESTED DOCUMENTS ({len(self.parsed_documents)} docs, {len(all_chunks)} chunks) ===\n\n"
        return header + "\n\n---\n\n".join(formatted)
    
    def clear(self):
        """Clear all ingested documents."""
        self.parsed_documents.clear()
        logger.info("Cleared all ingested documents")


# Convenience function for quick ingestion
def ingest_document(source: str, doc_type: Optional[str] = None) -> ParsedDocument:
    """Quick helper to ingest a single document."""
    engine = DocumentIngestionEngine()
    return engine.ingest(source, doc_type)
